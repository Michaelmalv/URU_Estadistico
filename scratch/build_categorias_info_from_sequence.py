import os
import re
import docx
import glob
import xml.etree.ElementTree as ET

# Rutas
DOC_DIR = 'documentos'
COMPONENTS_DIR = 'web/src/app/components'
OUTPUT_JS_PATH = os.path.join(COMPONENTS_DIR, 'categorias_info.js')

# Mapeo de categorías
MAPPING = {
    'Quitopía.docx': {
        'folder': 'quitopia',
        'key': 'Quitopia',
        'title': 'Quitopía'
    },
    'Recuperación Urbana Av. 10 de Agosto.docx': {
        'folder': 'recuperacion_urbana_av_10_de_agosto',
        'key': 'Recuperación Urbana Av. 10 de Agosto',
        'title': 'Recuperación Urbana Av. 10 de Agosto'
    },
    'REHABILITACIÓN DEL ESPACIO PÚBLICO.docx': {
        'folder': 'rehabilitacion_espacio_publico',
        'key': 'Rehabilitación del Espacio Público',
        'title': 'Rehabilitación del Espacio Público'
    },
    'Repotenciación Parque Bicentenario.docx': {
        'folder': 'repotenciacion_parque_bicentenario',
        'key': 'Repotenciación Parque Bicentenario',
        'title': 'Repotenciación Parque Bicentenario'
    },
    'Senderos Seguros.docx': {
        'folder': 'senderos_seguros',
        'key': 'Senderos Seguros',
        'title': 'Senderos Seguros'
    }
}

ORIGINAL_IMAGES = {
    'Quitopia': '/imagenes_categorias/quitopia/Quitopia.png',
    'Recuperación Urbana Av. 10 de Agosto': '/imagenes_categorias/recuperacion_urbana_av_10_de_agosto/recuperación_urbana_av10deAgosto.png',
    'Rehabilitación del Espacio Público': '/imagenes_categorias/rehabilitacion_espacio_publico/Rehabilitación del Centro Historico.jpg',
    'Repotenciación Parque Bicentenario': '/imagenes_categorias/repotenciacion_parque_bicentenario/Repotenciación del Bicentenario.png',
    'Senderos Seguros': '/imagenes_categorias/senderos_seguros/Senderos Seguros.png'
}

ORIGINAL_SUMMARIES = {
    'Quitopia': 'Red de centros de desarrollo comunitario y cuidado integral para el bienestar social de las familias.',
    'Recuperación Urbana Av. 10 de Agosto': 'Estrategia de revitalización del eje vial longitudinal de la 10 de Agosto para mejorar movilidad y vivienda.',
    'Rehabilitación del Espacio Público': 'Revitalización de calles patrimoniales para conectar peatonalmente y reactivar el comercio nocturno.',
    'Repotenciación Parque Bicentenario': 'Consolidación del espacio verde del antiguo aeropuerto mediante arborización nativa y nuevas áreas recreativas.',
    'Senderos Seguros': 'Corredores peatonales diseñados estratégicamente para mitigar índices delictivos y mejorar la seguridad de transeúntes.'
}

# Namespaces XML de Word
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
}

def clean_latex(text):
    text = re.sub(r'\[cite:\s*[^\]]+\]', '', text)
    text = re.sub(r'\$(\d+(?:[.,]\d+)?)\\text\{\s*[mM]\s*\}\^2\$', r'\1 m²', text)
    text = re.sub(r'\$(\d+(?:[.,]\d+)?)\\text\{\s*[hH]ect\w*\s*\}\$', r'\1 Hectáreas', text)
    text = re.sub(r'\\text\{\s*[mM]\s*\}\^2', 'm²', text)
    text = re.sub(r'\\text\{\s*([^}]+)\s*\}', r'\1', text)
    text = re.sub(r'USD\s*\$(\d+(?:[.,]\d+)?)\s*millones', r'USD \1 millones', text)
    text = re.sub(r'\$([^$]+)\$', r'\1', text)
    text = text.replace('\\', '')
    return text

def escape_jsx_braces(text):
    return text.replace('{', '&#123;').replace('}', '&#125;')

def parse_rels(doc_path):
    rels = {}
    import zipfile
    try:
        with zipfile.ZipFile(doc_path) as z:
            rels_data = z.read('word/_rels/document.xml.rels')
            root = ET.fromstring(rels_data)
            for rel in root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                r_id = rel.get('Id')
                target = rel.get('Target')
                if 'media/' in target:
                    rels[r_id] = os.path.basename(target)
    except Exception as e:
        print(f"Error parsing relationships for {doc_path}: {e}")
    return rels

def find_image_rids(p_element):
    rids = []
    for blip in p_element.findall('.//w:drawing//a:blip', NAMESPACES):
        rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if rid:
            rids.append(rid)
    for imagedata in p_element.findall('.//w:pict//v:imagedata', {'w': NAMESPACES['w'], 'v': 'urn:schemas-microsoft-com:vml'}):
        rid = imagedata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}href') or \
              imagedata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        if rid:
            rids.append(rid)
    return rids

def get_web_image_path(original_filename, folder_name):
    # original_filename is like 'image1.png'
    name, ext = os.path.splitext(original_filename)
    match = re.search(r'\d+', name)
    if match:
        idx = match.group(0)
        return f"/imagenes_categorias/{folder_name}/extracted_{idx}{ext}"
    return f"/imagenes_categorias/{folder_name}/{original_filename}"

def parse_runs_to_html(runs):
    html_parts = []
    for r in runs:
        txt = r.text
        if r.bold:
            lstrip = len(txt) - len(txt.lstrip())
            rstrip = len(txt) - len(txt.rstrip())
            sp_left = txt[:lstrip]
            sp_right = txt[len(txt)-rstrip:] if rstrip > 0 else ""
            core = txt.strip()
            if core:
                html_parts.append(f"{sp_left}<strong>{core}</strong>{sp_right}")
            else:
                html_parts.append(txt)
        else:
            html_parts.append(txt)
            
    html = "".join(html_parts)
    # Limpiar LaTeX al nivel de HTML
    # Buscar bloques de $ ... $ y limpiar su interior
    def replace_math(match):
        math_content = match.group(1)
        math_content_clean = re.sub(r'<[^>]+>', '', math_content)
        cleaned = clean_latex(f"${math_content_clean}$")
        return cleaned
    
    html = re.sub(r'\$([^$]+)\$', replace_math, html)
    html = clean_latex(html)
    html = escape_jsx_braces(html)
    return html

def build_jsx_for_docx(doc_path, info):
    doc = docx.Document(doc_path)
    rels = parse_rels(doc_path)
    
    jsx_blocks = []
    processed_paragraphs = set()
    
    # Buscamos subtítulo en el segundo párrafo
    sub_title = ""
    if len(doc.paragraphs) > 1:
        p1 = doc.paragraphs[1]
        if all(r.bold for r in p1.runs) if p1.runs else False:
            if p1.text.strip() != "Concepto del Proyecto":
                sub_title = escape_jsx_braces(clean_latex(p1.text.strip()))
                
    if sub_title:
        jsx_blocks.append(f'        <p style={{{{ color: "var(--text-muted)", fontSize: "1rem", fontWeight: "600", marginBottom: "1.5rem" }}}}>{sub_title}</p>')

    in_list = False
    list_type = "ul"
    
    # Para saltar el título y subtítulo iniciales
    skip_count = 2 if sub_title else 1
    
    # Recorrer elementos en secuencia
    body = doc.element.body
    p_index = 0
    
    for child in body:
        if child.tag.endswith('p'):
            p_obj = docx.text.paragraph.Paragraph(child, doc)
            text_raw = p_obj.text.strip()
            
            # Incrementar índice de párrafos de texto (excluyendo vacíos en el contador de saltos si es necesario)
            p_index += 1
            if p_index <= skip_count and text_raw:
                continue
                
            # Buscar imágenes en este párrafo
            rids = find_image_rids(child)
            images = [rels[rid] for rid in rids if rid in rels]
            
            if not text_raw:
                # Si el párrafo está vacío pero tiene imágenes
                if images:
                    if in_list:
                        jsx_blocks.append(f'        </{list_type}>')
                        in_list = False
                    for img in images:
                        web_path = get_web_image_path(img, info['folder'])
                        jsx_blocks.append(f'        <div className="detail-image-container">')
                        jsx_blocks.append(f'          <img src="{web_path}" alt="Detalle" />')
                        jsx_blocks.append(f'        </div>')
                continue
                
            runs = p_obj.runs
            is_bold = all(r.bold for r in runs) if runs else False
            is_heading = is_bold or text_raw.startswith(('1. ', '2. ', '3. ', '4. ', '5. ')) or text_raw == 'Concepto del Proyecto' or text_raw.startswith('Ficha Técnica')
            
            clean_text = escape_jsx_braces(clean_latex(text_raw))
            
            if is_heading:
                if in_list:
                    jsx_blocks.append(f'        </{list_type}>')
                    in_list = False
                    
                if text_raw == 'Concepto del Proyecto':
                    # Buscar el siguiente párrafo de texto
                    # En la estructura secuencial del XML
                    # Lo ponemos en la caja de alerta
                    # Para simplificar, buscamos en los siguientes hermanos
                    next_text_p = None
                    curr = child.getnext()
                    while curr is not None:
                        if curr.tag.endswith('p'):
                            temp_p = docx.text.paragraph.Paragraph(curr, doc)
                            if temp_p.text.strip():
                                next_text_p = temp_p
                                break
                        curr = curr.getnext()
                        
                    formatted_next = parse_runs_to_html(next_text_p.runs) if next_text_p else ""
                    
                    jsx_blocks.append('        <div className="info-alert">')
                    jsx_blocks.append(f'          {formatted_next}')
                    jsx_blocks.append('        </div>')
                    
                    # Si ese párrafo de concepto tenía imágenes asociadas, o el párrafo actual las tiene
                    concept_rids = find_image_rids(next_text_p._element) if next_text_p else []
                    concept_images = [rels[rid] for rid in concept_rids if rid in rels]
                    
                    all_concept_imgs = images + concept_images
                    for img in all_concept_imgs:
                        web_path = get_web_image_path(img, info['folder'])
                        jsx_blocks.append(f'        <div className="detail-image-container">')
                        jsx_blocks.append(f'          <img src="{web_path}" alt="Concepto" />')
                        jsx_blocks.append(f'        </div>')
                        
                    # El siguiente párrafo ya fue procesado, pero lo saltaremos en el bucle principal.
                    # Para hacer esto de forma limpia en el recorrido secuencial, marcamos que se debe ignorar cuando lleguemos a él.
                    # Guardamos el ID del elemento a ignorar
                    if next_text_p:
                        processed_paragraphs.add(next_text_p._element)
                else:
                    jsx_blocks.append(f'        <h3>{clean_text}</h3>')
                    # Renderizar imágenes del párrafo de cabecera
                    for img in images:
                        web_path = get_web_image_path(img, info['folder'])
                        jsx_blocks.append(f'        <div className="detail-image-container">')
                        jsx_blocks.append(f'          <img src="{web_path}" alt="Sección" />')
                        jsx_blocks.append(f'        </div>')
            else:
                # Comprobar si ya fue procesado como concepto
                if child in processed_paragraphs:
                    continue
                    
                has_list_pattern = False
                first_run_bold = runs[0].bold if runs else False
                first_run_text = runs[0].text.strip() if runs else ""
                
                if first_run_bold and (first_run_text.endswith(':') or (len(runs) > 1 and runs[1].text.strip().startswith(':'))):
                    has_list_pattern = True
                    
                if has_list_pattern:
                    if not in_list:
                        list_type = "ul"
                        jsx_blocks.append(f'        <{list_type}>')
                        in_list = True
                    item_html = parse_runs_to_html(runs)
                    jsx_blocks.append(f'          <li>{item_html}</li>')
                else:
                    if in_list:
                        jsx_blocks.append(f'        </{list_type}>')
                        in_list = False
                    p_html = parse_runs_to_html(runs)
                    jsx_blocks.append(f'        <p>{p_html}</p>')
                
                # Renderizar imágenes del párrafo
                for img in images:
                    if in_list:
                        jsx_blocks.append(f'        </{list_type}>')
                        in_list = False
                    web_path = get_web_image_path(img, info['folder'])
                    jsx_blocks.append(f'        <div className="detail-image-container">')
                    jsx_blocks.append(f'          <img src="{web_path}" alt="Detalle" />')
                    jsx_blocks.append(f'        </div>')
                    
        elif child.tag.endswith('tbl'):
            # Si hay una tabla
            if in_list:
                jsx_blocks.append(f'        </{list_type}>')
                in_list = False
                
            tbl_obj = docx.table.Table(child, doc)
            # Renderizar tabla de forma simple
            jsx_blocks.append('        <div className="table-responsive">')
            jsx_blocks.append('          <table className="table table-bordered">')
            jsx_blocks.append('            <tbody>')
            for row in tbl_obj.rows:
                jsx_blocks.append('              <tr>')
                for cell in row.cells:
                    cell_html = parse_runs_to_html(cell.paragraphs[0].runs) if cell.paragraphs else ""
                    jsx_blocks.append(f'                <td>{cell_html}</td>')
                jsx_blocks.append('              </tr>')
            jsx_blocks.append('            </tbody>')
            jsx_blocks.append('          </table>')
            jsx_blocks.append('        </div>')
            
            # Buscar imágenes en la tabla
            rids = find_image_rids(child)
            images = [rels[rid] for rid in rids if rid in rels]
            for img in images:
                web_path = get_web_image_path(img, info['folder'])
                jsx_blocks.append(f'        <div className="detail-image-container">')
                jsx_blocks.append(f'          <img src="{web_path}" alt="Tabla Detalle" />')
                jsx_blocks.append(f'        </div>')
                
    if in_list:
        jsx_blocks.append(f'        </{list_type}>')

    inner_content = "\n".join(jsx_blocks)
    
    # Obtener imagen de portada y resumen
    image_path = ORIGINAL_IMAGES.get(info['key'], f"/imagenes_categorias/{info['folder']}/{info['folder']}.png")
    summary_text = ORIGINAL_SUMMARIES.get(info['key'], "Detalles del proyecto...")
    
    return f"""  '{info['key']}': {{
    image: '{image_path}',
    summary: '{summary_text}',
    content: (
      <div className="article-content">
        <h1>{info['title']}</h1>
{inner_content}
      </div>
    )
  }}"""

# ==========================================
# DEFINICIONES HARDCODED (Zonas Metro, etc.)
# ==========================================

ZONAS_METRO_CODE = """  'Zonas Metro': {
    image: '/imagenes_categorias/zonas_metro/Zonas metro.jpg',
    summary: 'Modelo de ordenamiento físico-espacial en los exteriores de las estaciones para ordenar el flujo peatonal masivo.',
    content: (
      <div className="article-content">
        <h1>Zonas Metro</h1>
        <div className="info-alert">
          Para ordenar la alta concentración peatonal y el flujo masivo en los exteriores de las estaciones del Metro de Quito, se diseñó un <strong>modelo de ordenamiento físico-espacial</strong> que divide los accesos en tres zonas concéntricas:
        </div>
        <ul>
          <li><strong>Zona A - Aglomeración:</strong> Es el área colindante inmediata a la boca de ingreso del Metro. Está destinada de forma <strong>exclusiva</strong> a garantizar la entrada y salida fluida y segura de los pasajeros, por lo que debe mantenerse completamente libre de obstáculos.</li>
          <li><strong>Zona B - Concentración:</strong> Funciona como una franja o área de amortiguamiento urbano. En este espacio se permite la instalación regulada de <strong>mobiliario confortable</strong> (bancas, basureros, iluminación) diseñado para la permanencia o espera de los usuarios.</li>
          <li><strong>Zona C - Dispersión:</strong> Es el área abierta perimetral. Está perfilada como el lugar idóneo para la colocación de <strong>señalética de orientación</strong>, paradas de transporte complementario y casetas de <strong>comerciantes autónomos regularizados</strong> bajo el debido proceso de autorización.</li>
        </ul>

        <h3>Tótems de Identidad Metropolitana</h3>
        <p>Cada estación cuenta con un <strong>hito vertical unificado</strong> que le otorga identidad al entorno urbano. Este elemento técnico cuenta con las siguientes especificaciones y componentes:</p>
        <blockquote>
          <strong>Dimensiones estructurales:</strong> 4.04 metros de altura × 0.70 metros de ancho.
        </blockquote>
        <ul>
          <li><strong>Identificación:</strong> Incorpora el logo iconográfico de la estación y su nombre visible a larga distancia.</li>
          <li><strong>Orientación:</strong> Integra mapas detallados del sector y de la red completa del sistema de transporte.</li>
          <li><strong>Accesibilidad universal:</strong> Incluye un <strong>Mapa Háptico</strong> con texturas y relieves en braille especialmente diseñado para personas con discapacidad visual.</li>
          <li><strong>Información y servicios:</strong> Cuenta con espacios inferiores destinados a paneles informativos de la ciudad o publicidad regulada.</li>
        </ul>
      </div>
    )
  }"""

ARENA_BICENTENARIO_CODE = """  'Arena del Bicentenario': {
    image: '/imagenes_categorias/bicentenario/bicentenario.png',
    summary: 'Megaproyecto de renovación urbana de 105 hectáreas que integra humedales, ciclovías, deportes y una arena para conciertos.',
    content: (
      <div className="article-content">
        <h1>Arena del Bicentenario</h1>
        <div className="info-alert">
          La repotenciación de las 105 hectáreas del antiguo aeropuerto representa el proyecto de renovación y revitalización urbana más ambicioso del Distrito Metropolitano. Proyectado bajo las normativas del Plan de Uso y Gestión del Suelo (PUGS), este espacio técnico tiene el potencial de albergar vivienda y servicios para una población de hasta 304.000 habitantes, frenando la expansión descontrolada de las periferias.
        </div>

        <h3>Sostenibilidad e Infraestructura Verde-Azul</h3>
        <p>Con el fin de mitigar los efectos del cambio climático en la zona centro-norte, el parque implementa soluciones ambientales estratégicas:</p>
        <ul>
          <li><strong>Suelo permeable:</strong> Incremento del 5% de su superficie de absorción.</li>
          <li><strong>Red Verde-Azul:</strong> Consolidación de 8 hectáreas que contemplan la creación de <strong>6 cuerpos de agua y humedales</strong> diseñados para la captación técnica de aguas lluvias, la regulación del microclima y el fomento de la biodiversidad urbana.</li>
        </ul>

        <h3>Deporte, Comercio y Recreación Replicable</h3>
        <p>El megaproyecto organiza su espacio público a través de módulos replicables de equipamiento de alta calidad:</p>
        <ul>
          <li><strong>67 Canchas Deportivas:</strong> 21 de fútbol, 19 de voleibol, 11 de básquet, 7 de tenis, 5 múltiples, 3 de balonmano y 1 de béisbol.</li>
          <li><strong>Movilidad Activa:</strong> Construcción de <strong>12 km de caminerías internas</strong> iluminadas (divididas en 5 tipologías de diseño de estancia) y una red de <strong>14.4 km de ciclovías</strong> con carril segregado.</li>
          <li><strong>14 Plazas Comerciales:</strong> Equipadas con <strong>360 kioscos modulares</strong> de madera (1.80 × 1.80 m) integrados a zonas de vegetación y descanso para promover y ordenar el comercio local.</li>
          <li><strong>Zonas Temáticas:</strong> Incorporación de 12 áreas de juegos infantiles (con zonas especiales para la primera infancia), 38 zonas de picnic/descanso y 5 parques caninos (<em>dogparks</em>).</li>
        </ul>

        <h3>Grandes Hitos de Entretenimiento Masivo</h3>
        <ul>
          <li><strong>Arena de Espectáculos Quito:</strong> Megaestructura equipada con un escenario principal con capacidad para <strong>50.000 espectadores</strong>, diseñado para insertar a la capital en el circuito internacional de grandes conciertos, festivales y ferias masivas.</li>
          <li><strong>Anfiteatro Polifuncional:</strong> Espacio cultural al aire libre totalmente integrado al paisaje del parque, el cual utiliza relieves técnicos y graderíos naturales de césped para albergar eventos artísticos y comunitarios.</li>
        </ul>
      </div>
    )
  }"""

ARENA_DEL_SUR_CODE = """  'Arena del Sur': {
    image: '/imagenes_categorias/arena_del_sur/Arena del Sur.png',
    summary: 'Infraestructura cultural multifuncional en Quitumbe con centro cultural y plaza de espectáculos al aire libre.',
    content: (
      <div className="article-content">
        <h1>Quitumbe: Arena Cultural del Sur</h1>
        <div className="info-alert">
          Es un equipamiento cultural metropolitano de infraestructura inclusiva y multifuncional impulsado por la Alcaldía Metropolitana de Quito. Está diseñado para fortalecer la integración social, el acceso equitativo a la cultura y la dinamización comunitaria en el sur de la ciudad.
        </div>

        <h3>Cifras Clave del Proyecto</h3>
        <ul>
          <li><strong>Área de intervención:</strong> 33.944,72 m²</li>
          <li><strong>Inversión estimada:</strong> USD 6,2 millones</li>
          <li><strong>Ubicación:</strong> Centralidad Urbana Quitumbe</li>
        </ul>

        <h3>Componentes Principales (Plan de Intervención)</h3>
        <p>El proyecto articula sus actividades e infraestructura en dos grandes bloques de equipamiento:</p>
        <ol style={{ paddingLeft: '1.5rem', margin: '1rem 0' }}>
          <li style={{ marginBottom: '0.5rem' }}><strong>Centro Cultural Quitumbe:</strong> Un nodo de encuentro enfocado en la educación y las artes que incluye biblioteca, talleres artísticos, salas de danza, expresión corporal, auditorios y áreas educativas.</li>
          <li style={{ marginBottom: '0.5rem' }}><strong>Plaza de Espectáculos Quitumbe:</strong> Una infraestructura al aire libre dotada de una gran cubierta de diseño contemporáneo, optimizada para albergar eventos recreativos y conciertos masivos.</li>
        </ol>

        <h3>Propósito e Impacto</h3>
        <ul>
          <li><strong>Encuentro ciudadano:</strong> Funciona como un espacio para activar actividades artísticas, educativas y recreativas.</li>
          <li><strong>Fomento de la cohesión:</strong> Busca promover la apropiación pacífica del espacio público, la cohesión social y el bienestar de la población del sur de Quito.</li>
        </ul>
      </div>
    )
  }"""

# ==========================================
# PROCESAR Y GENERAR EN EL ORDEN DESEADO
# ==========================================

DESIRED_ORDER = [
    ('Senderos Seguros', 'Senderos Seguros.docx'),
    ('Zonas Metro', None),
    ('Rehabilitación del Espacio Público', 'REHABILITACIÓN DEL ESPACIO PÚBLICO.docx'),
    ('Repotenciación Parque Bicentenario', 'Repotenciación Parque Bicentenario.docx'),
    ('Quitopia', 'Quitopía.docx'),
    ('Recuperación Urbana Av. 10 de Agosto', 'Recuperación Urbana Av. 10 de Agosto.docx'),
    ('Arena del Bicentenario', None),
    ('Arena del Sur', None)
]

print("Generating sequential categories Javascript blocks in the desired order...")
generated_blocks = []

for key, doc_name in DESIRED_ORDER:
    if doc_name is not None:
        # Dynamic from docx
        doc_path = os.path.join(DOC_DIR, doc_name)
        if os.path.exists(doc_path):
            info = MAPPING[doc_name]
            print(f"Processing {doc_name} dynamically...")
            block = build_jsx_for_docx(doc_path, info)
            if block:
                generated_blocks.append(block)
        else:
            print(f"Warning: {doc_path} not found")
    else:
        # Hardcoded
        if key == 'Zonas Metro':
            generated_blocks.append(ZONAS_METRO_CODE)
        elif key == 'Arena del Bicentenario':
            generated_blocks.append(ARENA_BICENTENARIO_CODE)
        elif key == 'Arena del Sur':
            generated_blocks.append(ARENA_DEL_SUR_CODE)

# Unir todo
blocks_joined = ",\n\n".join(generated_blocks)
file_content = f"""'use strict';
import React from 'react';

export const CATEGORIAS_INFO = {{
{blocks_joined}
}};
"""

# Guardar en archivo de destino
os.makedirs(os.path.dirname(OUTPUT_JS_PATH), exist_ok=True)
with open(OUTPUT_JS_PATH, 'w', encoding='utf-8') as f:
    f.write(file_content)
    
print(f"Successfully generated {OUTPUT_JS_PATH}")
